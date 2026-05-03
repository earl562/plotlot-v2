"""Generate the PlotLot technical story as a formatted .docx file."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def set_run_font(run, size=11, bold=False, italic=False, color=None, name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_styled_paragraph(doc, text, size=11, bold=False, italic=False, color=None, alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if alignment:
        pf.alignment = alignment
    return p


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_mixed_paragraph(doc, segments, space_after=6):
    """Add a paragraph with mixed formatting. segments is a list of (text, kwargs) tuples."""
    p = doc.add_paragraph()
    for text, kwargs in segments:
        run = p.add_run(text)
        set_run_font(run, **kwargs)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, name="Courier New", color=(40, 40, 40))
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.4)
    pf.line_spacing = Pt(13)
    return p


def add_callout(doc, text):
    """Add an emphasized callout/insight box."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, italic=True, color=(0, 100, 160))
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(16)
    return p


def add_section_heading(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(30, 60, 110)
        run.font.size = Pt(18)
    return h


def add_subsection(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(50, 90, 140)
        run.font.size = Pt(14)
    return h


def build_document():
    doc = Document()

    # -- Default style adjustments --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    add_styled_paragraph(
        doc, "How I Built PlotLot",
        size=28, bold=True, color=(30, 60, 110),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
    )
    add_styled_paragraph(
        doc, "An AI System That Reads Zoning Law\nSo Developers Don't Have To",
        size=16, italic=True, color=(80, 80, 80),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24,
    )
    add_styled_paragraph(
        doc, "Earl Perry",
        size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    add_styled_paragraph(
        doc, "ML/LLMOps Engineer",
        size=12, italic=True, color=(100, 100, 100),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    add_styled_paragraph(
        doc, "February 2026",
        size=12, color=(100, 100, 100),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS (manual)
    # ============================================================
    add_styled_paragraph(doc, "Table of Contents", size=18, bold=True, color=(30, 60, 110), space_after=12)

    toc_items = [
        ("1.", "The Problem That Started Everything"),
        ("2.", "Getting the Data \u2014 The Part Nobody Talks About"),
        ("3.", "Search \u2014 Why Vector Search Alone Isn't Enough"),
        ("4.", "The Pipeline \u2014 Where the Agent Lives"),
        ("5.", "Evaluation \u2014 Because 'It Seems To Work' Isn't Good Enough"),
        ("6.", "Observability \u2014 What Happens at 3am"),
        ("7.", "Infrastructure \u2014 Making It Real"),
        ("8.", "What I'd Do Differently"),
        ("9.", "The Numbers"),
    ]
    for num, title in toc_items:
        add_mixed_paragraph(doc, [
            (f"  {num}  ", {"size": 11, "bold": True, "color": (30, 60, 110)}),
            (title, {"size": 11}),
        ], space_after=4)

    doc.add_page_break()

    # ============================================================
    # CHAPTER 1 — THE PROBLEM
    # ============================================================
    add_section_heading(doc, "1. The Problem That Started Everything")

    add_body(doc, "I was looking at real estate in South Florida and hit a wall that every developer, investor, and architect hits: zoning.")

    add_body(doc, "You find a property. You want to know: can I build a duplex? A fourplex? How tall can I go? What are the setbacks? The answers exist \u2014 buried in municipal zoning ordinances that read like they were written by lawyers who hate you. And here's the kicker: South Florida has 104 municipalities across three counties. Miami-Dade alone has 35. Each one has its own zoning code, its own terminology, its own way of defining \"residential density.\"")

    add_body(doc, "A human doing this manually is looking at hours of research per property. A real estate developer with a portfolio is paying a land use attorney $500/hour to do it.")

    add_callout(doc, "I thought: this is a retrieval problem, an extraction problem, and a math problem. I can automate all three.")

    doc.add_page_break()

    # ============================================================
    # CHAPTER 2 — DATA INGESTION
    # ============================================================
    add_section_heading(doc, "2. Getting the Data \u2014 The Part Nobody Talks About")

    add_body(doc, "The first thing I learned is that the glamorous part of AI \u2014 the model, the prompts, the agent \u2014 is maybe 20% of the work. The other 80% is getting clean data into a place where you can search it.")

    add_body(doc, "All 104 municipalities publish their zoning ordinances on Municode. That's the good news. The bad news is there's no standard. Miami Gardens calls it \"Chapter 24 \u2014 Zoning.\" Fort Lauderdale calls it \"Unified Land Development Regulations.\" Coral Gables buries it under \"City Code, Part III, Article 3.\"")

    add_subsection(doc, "Auto-Discovery: Let the Machine Find the Data")

    add_body(doc, "I wasn't going to hardcode 104 municipality configs and maintain them by hand. So I built an auto-discovery system. It hits the Municode Library API, searches by county, walks each municipality's table of contents tree, and uses heuristics to find the zoning chapter \u2014 looking for titles containing \"zoning,\" \"land development,\" or \"land use.\"")

    add_body(doc, "First time I ran it, it found 73 out of 104. The rest either don't use Municode or have non-standard structures I'll handle later. 73 municipalities covering the vast majority of South Florida? That's a strong starting point.")

    add_subsection(doc, "Scraping, Chunking, Embedding")

    add_body(doc, "Once I know where the zoning chapter is, the scraper pulls every section. Raw HTML comes back. The chunker strips the markup, splits on section boundaries, and attaches metadata to every chunk \u2014 which municipality, which county, which chapter, which section, and critically, which zone codes are mentioned in the text. That metadata is what makes retrieval precise later.")

    add_body(doc, "Then embeddings. I'm using HuggingFace's Inference API with a 1024-dimensional model. Batch processing, retry logic, nothing exotic.")

    add_subsection(doc, "The Bug That Almost Killed Search Quality")

    add_body(doc, "During early testing, I noticed search quality was randomly terrible for certain municipalities. Chunks that should have been top results were nowhere. I dug in and found the culprit: zero vectors.")

    add_body(doc, "The embedding API was occasionally returning all-zeros on timeout \u2014 no error, no exception, just a silent zero vector. And a zero vector in cosine similarity space has maximum similarity to everything. It was polluting my results silently.")

    add_body(doc, "So I added a validation step before anything touches the database:")

    add_code_block(doc, (
        "def validate_chunks(chunks, embeddings):\n"
        "    for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):\n"
        "        if len(emb) != EMBEDDING_DIM:        # wrong dimension\n"
        "        if all(v == 0.0 for v in emb):        # zero vector\n"
        "        if len(chunk.text.strip()) < 50:      # degenerate chunk"
    ))

    add_callout(doc, "In production ML systems, data validation at ingestion time is the single highest-ROI investment you can make. Most teams learn this the hard way after months of mysterious quality degradation.")

    doc.add_page_break()

    # ============================================================
    # CHAPTER 3 — SEARCH / RETRIEVAL
    # ============================================================
    add_section_heading(doc, "3. Search \u2014 Why Vector Search Alone Isn't Enough")

    add_body(doc, "With the data in place, I needed retrieval. The obvious approach is vector search \u2014 embed the query, find the nearest neighbors. I tried it. It mostly worked. Then it didn't.")

    add_body(doc, "The query \"R-1 residential setback requirements Miami Gardens\" was returning chunks about R-2 zoning in Miramar. Why? Because semantically, those concepts are very similar. They're both residential zoning setback discussions. The embedding model can't distinguish \"R-1 in Miami Gardens\" from \"R-2 in Miramar\" because the meaning is almost identical. Only the exact terms differ.")

    add_callout(doc, "Dense embeddings capture semantics but lose lexical precision. This is a well-known problem in information retrieval.")

    add_subsection(doc, "Hybrid Search with Reciprocal Rank Fusion")

    add_body(doc, "I run two searches in parallel against the same PostgreSQL database:")

    add_mixed_paragraph(doc, [
        ("1. Vector search", {"size": 11, "bold": True}),
        (" \u2014 pgvector's cosine distance against the query embedding", {"size": 11}),
    ])
    add_mixed_paragraph(doc, [
        ("2. Full-text BM25 search", {"size": 11, "bold": True}),
        (" \u2014 PostgreSQL's built-in tsvector and ts_rank", {"size": 11}),
    ])

    add_body(doc, "Vector search finds conceptually relevant chunks. Text search finds chunks that contain the exact zone code \"R-1\" and the exact word \"Miami Gardens.\" Then I combine them with Reciprocal Rank Fusion:")

    add_code_block(doc, "RRF_score(doc) = 1/(k + rank_vector) + 1/(k + rank_text)")

    add_body(doc, "With k=60. A document that ranks well in both searches scores much higher than one that ranks well in only one. The exact zone code match from BM25 anchors the results, and the semantic similarity from vectors brings in related context.")

    add_subsection(doc, "Why RRF Over a Learned Ranker?")

    add_body(doc, "Two reasons. First, I don't have click-through data for training a learned model. RRF is unsupervised. Second, RRF is deterministic and interpretable. When a result is wrong, I can look at the vector rank and the BM25 rank independently and understand why the fusion produced that result. With a learned model, it's a black box.")

    add_body(doc, "Also \u2014 before search even runs, I filter by municipality. This isn't a soft signal in the ranking; it's a hard WHERE clause. If you're querying about Miami Gardens, you should never see Fort Lauderdale results regardless of semantic similarity. I've seen RAG systems in production that treat metadata filtering as optional. It's not. It's the difference between a useful system and a hallucination factory.")

    doc.add_page_break()

    # ============================================================
    # CHAPTER 4 — THE PIPELINE
    # ============================================================
    add_section_heading(doc, "4. The Pipeline \u2014 Where the Agent Lives")

    add_body(doc, "Now I have good data and good retrieval. Time to build the actual product: address in, investment analysis out.")

    add_body(doc, "The pipeline has five stages. Each one feeds the next.")

    # Stage table
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    headers = ["Stage", "What It Does", "Key Detail"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)

    rows_data = [
        ("Geocoding", "Address \u2192 county + municipality", "Geocodio API; municipality is the pivot"),
        ("Property Lookup", "Coordinates \u2192 parcel data", "3 county ArcGIS APIs, each different"),
        ("Hybrid Retrieval", "Zone code \u2192 relevant ordinance chunks", "RRF fusion, hard municipality filter"),
        ("LLM Extraction", "Ordinance text \u2192 numeric parameters", "Tool calling, not free-text output"),
        ("Calculator", "Parameters + lot area \u2192 max units", "Pure math, no LLM, deterministic"),
    ]
    for row_idx, (stage, what, detail) in enumerate(rows_data, start=1):
        for col_idx, val in enumerate([stage, what, detail]):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    add_styled_paragraph(doc, "", space_after=6)  # spacer

    add_subsection(doc, "The County API Headache")

    add_body(doc, "Property lookup was more painful than I expected. Each of the three counties has a different ArcGIS REST API. Miami-Dade is actually two API calls \u2014 one layer for parcel data (folio number, lot size), a completely separate layer for zoning designation. Broward combines them but uses different field names. Palm Beach is different again.")

    add_body(doc, "I built an adapter for each county behind a common interface. The pipeline doesn't know or care which county it's talking to \u2014 it gets back a PropertyInfo Pydantic model with the zone code, lot area, and dimensions.")

    add_subsection(doc, "The Design Decision I'm Most Proud Of")

    add_body(doc, "The tempting approach is: give the LLM the zoning text and ask \"how many units can I build?\" Let it do the whole thing. But I've seen what happens when LLMs do math. They hallucinate numbers. They mix up units. They'll say \"density is 12 units per acre\" when the ordinance says \"minimum lot area is 12,000 square feet per unit\" \u2014 a completely different number that leads to a completely different answer.")

    add_callout(doc, "So I split it. The LLM's job is extraction only. It reads text and fills in numbers. Math does the rest.")

    add_body(doc, "I give the LLM a structured tool definition \u2014 a Pydantic schema called NumericZoningParams:")

    add_code_block(doc, (
        "class NumericZoningParams(BaseModel):\n"
        "    max_density_units_per_acre: float | None\n"
        "    min_lot_area_sq_ft: float | None\n"
        "    max_far: float | None\n"
        "    max_lot_coverage_pct: float | None\n"
        "    max_height_ft: float | None\n"
        "    min_front_setback_ft: float | None\n"
        "    min_side_setback_ft: float | None\n"
        "    min_rear_setback_ft: float | None"
    ))

    add_body(doc, "The LLM reads the ordinance text and calls this tool with the numbers it finds. If a value isn't in the text, it returns None \u2014 not a guess. The prompt is explicit: \"Only extract values explicitly stated in the provided text. Do not infer, calculate, or estimate.\"")

    add_subsection(doc, "Why Tool Calling Instead of JSON Output?")

    add_body(doc, "Two reasons. First, tool calling gives you schema validation for free \u2014 if the model returns a string where a float should be, the framework catches it before my code ever sees it. Second, it's model-agnostic. I run Kimi K2.5 on NVIDIA NIM as primary, DeepSeek V3.2 as fallback. Both support the same tool calling interface.")

    add_subsection(doc, "The Calculator: Where Math Happens")

    add_body(doc, "Pure Python, no LLM, no ambiguity. It takes the extracted parameters and the lot area, and applies every constraint independently:")

    add_code_block(doc, (
        "# Density constraint\n"
        "units = (lot_area / 43560) * units_per_acre\n\n"
        "# Lot area constraint\n"
        "units = lot_area / min_lot_per_unit\n\n"
        "# FAR constraint\n"
        "buildable = lot_area * FAR\n"
        "units = buildable / avg_unit_size\n\n"
        "# The binding constraint is the MINIMUM\n"
        "max_units = min(c.units for c in constraints)"
    ))

    add_body(doc, "The output includes a full breakdown \u2014 which constraint limits you, what each constraint allows, and a confidence level. HIGH if multiple constraints agree, MEDIUM if only one applies, LOW if we're missing key parameters.")

    add_callout(doc, "This separation \u2014 LLM extracts, math decides \u2014 makes the system testable, debuggable, and trustworthy. It's the single most important architectural decision in the system.")

    doc.add_page_break()

    # ============================================================
    # CHAPTER 5 — EVALUATION
    # ============================================================
    add_section_heading(doc, "5. Evaluation \u2014 Because 'It Seems To Work' Isn't Good Enough")

    add_body(doc, "I have a pipeline that works on the examples I've tested. But how do I know it works on the municipalities I haven't tested? And how do I know a prompt change doesn't break the ones that already work?")

    add_subsection(doc, "The Golden Dataset")

    add_body(doc, "Curated examples with known correct answers. \"Miami Gardens, R-1, 7,500 sq ft lot \u2192 max 1 unit, density is binding, confidence HIGH.\" These are manually verified against the actual ordinance text.")

    add_subsection(doc, "Deterministic Scorers (Not LLM-as-Judge)")

    add_body(doc, "I built rule-based scorers, not LLM-as-judge. Using an LLM to evaluate an LLM is circular.")

    add_mixed_paragraph(doc, [
        ("Numeric extraction accuracy", {"size": 11, "bold": True}),
        (" \u2014 Did the LLM extract the right density value? Exact match, not fuzzy.", {"size": 11}),
    ])
    add_mixed_paragraph(doc, [
        ("Unit detection", {"size": 11, "bold": True}),
        (" \u2014 Did we get the right unit (units/acre vs sq ft/unit)? Getting the number right but the unit wrong is worse than getting nothing.", {"size": 11}),
    ])
    add_mixed_paragraph(doc, [
        ("Report completeness", {"size": 11, "bold": True}),
        (" \u2014 Did the output include all expected fields?", {"size": 11}),
    ])

    add_body(doc, "These scorers run through mlflow.genai.evaluate(). Every eval run is an MLflow experiment with logged metrics, artifacts, and the exact prompt version used. When I change the extraction prompt, the next eval run shows me: \"Prompt v3 got 82% extraction accuracy, Prompt v4 got 91%.\" All in the MLflow UI, all traceable, all comparable.")

    add_subsection(doc, "Quality Gates That Block Bad Merges")

    add_code_block(doc, (
        'DEFAULT_THRESHOLDS = {\n'
        '    "report_completeness/mean": 0.6,\n'
        '}'
    ))

    add_body(doc, "If any metric drops below its threshold, the pipeline fails. In CI, a quality gate script returns a non-zero exit code, which blocks the merge. No human has to remember to check quality. The system enforces it.")

    add_callout(doc, "This is the pattern that separates experimental ML from production ML. Experimentation is \"I tried it, it looks good.\" Production is \"I have automated regression detection that blocks bad changes before they reach users.\"")

    doc.add_page_break()

    # ============================================================
    # CHAPTER 6 — OBSERVABILITY
    # ============================================================
    add_section_heading(doc, "6. Observability \u2014 What Happens at 3am")

    add_body(doc, "The pipeline works. The eval catches regressions. But what happens when the system is live and something goes wrong?")

    add_subsection(doc, "Structured JSON Logging")

    add_body(doc, "Every log line is a JSON object:")

    add_code_block(doc, (
        '{\n'
        '  "timestamp": "2026-02-17T15:30:00+00:00",\n'
        '  "level": "INFO",\n'
        '  "logger": "plotlot.retrieval.search",\n'
        '  "message": "Hybrid search returned 8 chunks",\n'
        '  "correlation_id": "abc-123-def",\n'
        '  "municipality": "miami-gardens",\n'
        '  "duration_ms": 45\n'
        '}'
    ))

    add_body(doc, "Not print() statements. Not unstructured text that you grep through with regex prayers. Structured fields that Datadog, Grafana Loki, or any log aggregator can index and query.")

    add_subsection(doc, "Correlation IDs: One ID, One Request, Every Log Line")

    add_body(doc, "The correlation ID is a ContextVar \u2014 Python's async-safe equivalent of thread-local storage. When a request hits the FastAPI server, middleware generates a UUID (or uses the X-Request-ID header if the load balancer set one) and stores it. From that point on, every await in the chain \u2014 geocoding, property lookup, search, LLM call, calculator \u2014 every log line automatically includes that ID.")

    add_code_block(doc, (
        'correlation_id: ContextVar[str] = ContextVar("correlation_id", default="")\n\n'
        "class CorrelationIDMiddleware(BaseHTTPMiddleware):\n"
        "    async def dispatch(self, request, call_next):\n"
        '        cid = request.headers.get("x-request-id", str(uuid.uuid4()))\n'
        "        token = correlation_id.set(cid)\n"
        "        response = await call_next(request)\n"
        '        response.headers["x-request-id"] = cid\n'
        "        return response"
    ))

    add_body(doc, "The alternative is passing the correlation ID as a parameter through every function in the entire codebase. ContextVar is implicit propagation \u2014 any function anywhere calls get_correlation_id() and gets the right value for the current async context. The trade-off is explicitness versus ergonomics, and in a deep async call stack, ergonomics wins.")

    add_subsection(doc, "MLflow Tracing: The Full Picture")

    add_body(doc, "Every pipeline run is an MLflow trace with spans for each stage. I can open the MLflow UI and see: this request took 2.3 seconds. 150ms was geocoding, 200ms was property lookup, 180ms was retrieval, 1.8 seconds was the LLM call. Token counts, similarity scores, calculator outputs \u2014 all logged.")

    add_body(doc, "This isn't just debugging \u2014 it's optimization intelligence. When I see that the LLM call is 78% of total latency, I know where to focus. When retrieval scores are consistently low for Palm Beach municipalities, I know my Palm Beach ingestion needs attention.")

    add_subsection(doc, "The Health Check")

    add_code_block(doc, (
        'GET /health \u2192 {\n'
        '  "status": "healthy",\n'
        '  "checks": {\n'
        '    "database": "ok",\n'
        '    "last_ingestion": "2026-02-15T02:00:00+00:00",\n'
        '    "mlflow": "ok"\n'
        '  }\n'
        '}'
    ))

    add_body(doc, "If the database is down, status is \"degraded.\" If the last ingestion was 10 days ago, that's a signal the weekly Prefect flow failed. If MLflow is unreachable, tracing is broken but the API still serves. This is what an on-call engineer looks at on their phone at 3am to decide: do I need to get out of bed?")

    doc.add_page_break()

    # ============================================================
    # CHAPTER 7 — INFRASTRUCTURE
    # ============================================================
    add_section_heading(doc, "7. Infrastructure \u2014 Making It Real")

    add_body(doc, "A system that only runs on my laptop isn't a system. It's a demo.")

    add_subsection(doc, "Multi-Target Dockerfile")

    add_body(doc, "One Dockerfile, three build targets. The runtime stage has the Python environment. Then:")

    add_code_block(doc, (
        "FROM runtime AS api       # FastAPI server\n"
        "FROM runtime AS ingest    # One-shot ingestion job\n"
        "FROM runtime AS worker    # Prefect worker for scheduled flows"
    ))

    add_body(doc, "docker build --target api gives you the web server. --target ingest gives you the data pipeline. --target worker gives you the orchestration runner. Same Python environment, same dependency versions, different entrypoints. No drift.")

    add_subsection(doc, "Five-Service Docker Compose Stack")

    table2 = doc.add_table(rows=6, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = "Light Grid Accent 1"

    for i, h in enumerate(["Service", "Purpose"]):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)

    services = [
        ("db", "PostgreSQL + pgvector \u2014 chunks, embeddings, relational data"),
        ("api", "FastAPI application \u2014 handles user requests"),
        ("mlflow", "MLflow tracking server \u2014 experiments, tracing, registry"),
        ("prefect-server", "Orchestration UI + API \u2014 flow scheduling and monitoring"),
        ("prefect-worker", "Executes scheduled flows \u2014 ingestion and eval jobs"),
    ]
    for row_idx, (svc, purpose) in enumerate(services, start=1):
        table2.rows[row_idx].cells[0].text = svc
        table2.rows[row_idx].cells[1].text = purpose
        for col_idx in range(2):
            for p in table2.rows[row_idx].cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    add_styled_paragraph(doc, "", space_after=6)

    add_body(doc, "One docker compose up command. Full production-equivalent stack on your machine.")

    add_subsection(doc, "Prefect Orchestration")

    add_body(doc, "Two scheduled deployments keep the system healthy without human intervention:")

    add_mixed_paragraph(doc, [
        ("Weekly ingestion", {"size": 11, "bold": True}),
        (" \u2014 Every Sunday at 2am, re-scrape all 73 municipalities. Catches ordinance amendments.", {"size": 11}),
    ])
    add_mixed_paragraph(doc, [
        ("Nightly eval", {"size": 11, "bold": True}),
        (" \u2014 Every night at 4am, run the golden dataset through all scorers. Catches quality regressions before anyone notices.", {"size": 11}),
    ])

    add_body(doc, "The Prefect decorators are designed with graceful fallback. If Prefect isn't installed \u2014 like in CI \u2014 the functions run as plain async Python. Same code, same behavior, just no orchestration metadata. No separate \"test mode\" and \"production mode\" code paths.")

    add_subsection(doc, "CI/CD: Two GitHub Actions Workflows")

    add_mixed_paragraph(doc, [
        ("CI workflow", {"size": 11, "bold": True}),
        (" \u2014 Every push: ruff lint, ruff format check, pytest against a real pgvector service container. Not mocked Postgres \u2014 actual PostgreSQL with the pgvector extension.", {"size": 11}),
    ])
    add_mixed_paragraph(doc, [
        ("Eval workflow", {"size": 11, "bold": True}),
        (" \u2014 Main branch pushes: validates the evaluation framework itself. Meta-testing \u2014 testing the tests.", {"size": 11}),
    ])

    add_body(doc, "A quality gate script checks MLflow metrics against thresholds. If extraction accuracy drops below the threshold, the script exits non-zero, the CI step fails, the merge is blocked. Automated quality enforcement.")

    doc.add_page_break()

    # ============================================================
    # CHAPTER 8 — WHAT I'D DO DIFFERENTLY
    # ============================================================
    add_section_heading(doc, "8. What I'd Do Differently")

    add_body(doc, "If I started over, three things:")

    add_mixed_paragraph(doc, [
        ("1. Schema versioning from day one. ", {"size": 11, "bold": True}),
        ("When I added created_at to the chunks table, I did it by modifying the SQLAlchemy model. In production, that's an Alembic migration. I haven't set up Alembic yet. I should have.", {"size": 11}),
    ])

    add_mixed_paragraph(doc, [
        ("2. Async embedding batches. ", {"size": 11, "bold": True}),
        ("The embedding step is synchronous and sequential. At 73 municipalities with hundreds of chunks each, this is the slowest part of ingestion. Async batching with concurrency limits would cut ingestion time significantly.", {"size": 11}),
    ])

    add_mixed_paragraph(doc, [
        ("3. More golden data sooner. ", {"size": 11, "bold": True}),
        ("73 municipalities and I have golden examples for a handful. Every municipality I add to the golden dataset is another regression I'll catch automatically. Building the eval framework was the right call; I should have populated it more aggressively.", {"size": 11}),
    ])

    doc.add_page_break()

    # ============================================================
    # CHAPTER 9 — THE NUMBERS
    # ============================================================
    add_section_heading(doc, "9. The Numbers")

    stats = [
        ("73", "municipalities auto-discovered and ingested"),
        ("3", "counties with distinct API adapters"),
        ("261", "tests \u2014 242 unit, 12 scorer, 7 eval"),
        ("5", "services in the Docker Compose stack"),
        ("100%", "pipeline runs traced in MLflow"),
        ("Every", "log line structured JSON with correlation IDs"),
        ("Automated", "quality gates blocking deployment on regression"),
        ("Zero", "hardcoded municipality configs for the core 73"),
    ]

    table3 = doc.add_table(rows=len(stats), cols=2)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, (num, desc) in enumerate(stats):
        cell_num = table3.rows[row_idx].cells[0]
        cell_desc = table3.rows[row_idx].cells[1]
        p_num = cell_num.paragraphs[0]
        r_num = p_num.add_run(num)
        set_run_font(r_num, size=14, bold=True, color=(30, 60, 110))
        p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p_desc = cell_desc.paragraphs[0]
        r_desc = p_desc.add_run(desc)
        set_run_font(r_desc, size=11)

    add_styled_paragraph(doc, "", space_after=12)

    # ============================================================
    # CLOSING
    # ============================================================
    add_styled_paragraph(doc, "", space_after=24)

    add_styled_paragraph(
        doc,
        "The system reads law so humans don't have to.",
        size=14, bold=True, italic=True, color=(30, 60, 110),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
    )
    add_styled_paragraph(
        doc,
        "It tells you what you can build, shows you which constraint limits you,\nand gives you a confidence level on the answer.",
        size=12, italic=True, color=(80, 80, 80),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
    )
    add_styled_paragraph(
        doc,
        "And every piece of it \u2014 from ingestion validation to the eval quality gate \u2014\nis designed so that when something breaks, we know before the user does.",
        size=12, italic=True, color=(80, 80, 80),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24,
    )
    add_styled_paragraph(
        doc,
        "That's PlotLot.",
        size=16, bold=True, color=(30, 60, 110),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_path = Path(__file__).resolve().parent.parent / "How_I_Built_PlotLot.docx"
    doc.save(str(output_path))
    print(f"Document saved to: {output_path}")
