"""Contract and deal document generation (DEPRECATED).

This module is superseded by the Clause Builder system in
``plotlot.clauses``. The ``generate_loi`` and ``generate_deal_summary``
functions are kept for backward compatibility but internally route
through the clause engine.

Canonical location for ``GeneratedDocument`` is now
``plotlot.clauses.schema.GeneratedDocument``.
"""

from __future__ import annotations

import io
import logging
import warnings
from dataclasses import dataclass, field
from datetime import datetime, timezone

from plotlot.clauses.schema import GeneratedDocument  # canonical import  # noqa: F401
from plotlot.core.types import (
    LandProForma,
    PropertyRecord,
    ZoningReport,
)

logger = logging.getLogger(__name__)


@dataclass
class LOIParams:
    """Parameters for Letter of Intent generation."""

    buyer_name: str = ""
    buyer_entity: str = ""
    buyer_email: str = ""
    buyer_phone: str = ""
    earnest_money_pct: float = 1.0  # % of offer price
    due_diligence_days: int = 30
    closing_days: int = 60
    contingencies: list[str] = field(
        default_factory=lambda: [
            "Satisfactory zoning verification",
            "Environmental assessment (Phase I)",
            "Clear title and survey",
        ]
    )


def generate_loi(
    report: ZoningReport,
    pro_forma: LandProForma | None = None,
    params: LOIParams | None = None,
) -> GeneratedDocument:
    """Generate a Letter of Intent .docx from analysis results.

    .. deprecated::
        Use ``plotlot.clauses.engine.assemble_document()`` with
        ``DocumentType.loi`` instead.
    """
    warnings.warn(
        "generate_loi() is deprecated. Use plotlot.clauses.engine.assemble_document() "
        "with DocumentType.loi instead.",
        DeprecationWarning,
        stacklevel=2,
    )
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx not installed — generating plain text LOI")
        return _generate_text_loi(report, pro_forma, params)

    if params is None:
        params = LOIParams()

    offer_price = pro_forma.max_land_price if pro_forma else 0
    earnest_money = offer_price * (params.earnest_money_pct / 100)
    now = datetime.now(timezone.utc)
    prop = report.property_record or PropertyRecord()

    doc = Document()

    # Title
    title = doc.add_heading("Letter of Intent", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    doc.add_paragraph(f"Date: {now.strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    # Subject
    doc.add_heading("RE: Letter of Intent to Purchase", level=2)

    # Property description
    doc.add_paragraph(
        f"This Letter of Intent outlines the terms under which "
        f"{params.buyer_name or '[BUYER NAME]'} "
        f"({params.buyer_entity or '[BUYER ENTITY]'}) proposes to purchase "
        f"the property located at:"
    )
    doc.add_paragraph(f"Address: {report.formatted_address or report.address}")
    doc.add_paragraph(f"Folio/Parcel: {prop.folio or 'N/A'}")
    doc.add_paragraph(f"Lot Size: {prop.lot_size_sqft:,.0f} sqft")
    doc.add_paragraph(f"Zoning: {report.zoning_district}")
    doc.add_paragraph("")

    # Terms table
    doc.add_heading("Proposed Terms", level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"

    terms = [
        ("Purchase Price", f"${offer_price:,.0f}" if offer_price > 0 else "[TO BE DETERMINED]"),
        ("Earnest Money", f"${earnest_money:,.0f}" if earnest_money > 0 else "[TBD]"),
        ("Due Diligence Period", f"{params.due_diligence_days} days from execution"),
        ("Closing Date", f"{params.closing_days} days from execution"),
        ("Financing", "Cash or conventional financing"),
        ("Title", "Marketable title, free of liens and encumbrances"),
        ("Property Condition", "As-is, where-is"),
    ]
    for i, (label, value) in enumerate(terms):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph("")

    # Contingencies
    doc.add_heading("Contingencies", level=2)
    for contingency in params.contingencies:
        doc.add_paragraph(contingency, style="List Bullet")
    doc.add_paragraph("")

    # Development potential
    doc.add_heading("Development Potential Summary", level=2)
    if report.density_analysis:
        doc.add_paragraph(
            f"Maximum Allowable Units: {report.density_analysis.max_units} "
            f"({report.density_analysis.governing_constraint})"
        )
    if pro_forma and pro_forma.gross_development_value > 0:
        doc.add_paragraph(f"Gross Development Value: ${pro_forma.gross_development_value:,.0f}")
        doc.add_paragraph(f"Estimated Hard Costs: ${pro_forma.hard_costs:,.0f}")
        doc.add_paragraph(f"Maximum Land Price (Residual): ${pro_forma.max_land_price:,.0f}")
    doc.add_paragraph("")

    # Signature
    doc.add_heading("Acceptance", level=2)
    doc.add_paragraph(
        "This LOI is non-binding and subject to execution of a formal Purchase and Sale Agreement."
    )
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph(params.buyer_name or "[BUYER NAME]")
    doc.add_paragraph(params.buyer_entity or "[BUYER ENTITY]")
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Seller Signature")
    doc.add_paragraph("")

    # Footer
    p = doc.add_paragraph()
    p.add_run("Generated by PlotLot — ").italic = True
    p.add_run(now.strftime("%Y-%m-%d %H:%M UTC")).italic = True

    buf = io.BytesIO()
    doc.save(buf)

    address_slug = (report.address or "property").split(",")[0].replace(" ", "_")[:30]
    return GeneratedDocument(
        filename=f"LOI_{address_slug}_{now.strftime('%Y%m%d')}.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buf.getvalue(),
    )


def _generate_text_loi(
    report: ZoningReport,
    pro_forma: LandProForma | None,
    params: LOIParams | None,
) -> GeneratedDocument:
    """Fallback plain-text LOI when python-docx is not available."""
    if params is None:
        params = LOIParams()

    offer_price = pro_forma.max_land_price if pro_forma else 0
    now = datetime.now(timezone.utc)
    prop = report.property_record or PropertyRecord()

    lines = [
        "LETTER OF INTENT",
        f"Date: {now.strftime('%B %d, %Y')}",
        "",
        f"Buyer: {params.buyer_name or '[BUYER NAME]'}",
        f"Entity: {params.buyer_entity or '[BUYER ENTITY]'}",
        "",
        f"Property: {report.formatted_address or report.address}",
        f"Folio: {prop.folio or 'N/A'}",
        f"Lot Size: {prop.lot_size_sqft:,.0f} sqft",
        f"Zoning: {report.zoning_district}",
        "",
        f"Purchase Price: ${offer_price:,.0f}" if offer_price > 0 else "Purchase Price: [TBD]",
        f"Due Diligence: {params.due_diligence_days} days",
        f"Closing: {params.closing_days} days",
        "",
        f"Generated by PlotLot — {now.strftime('%Y-%m-%d %H:%M UTC')}",
    ]

    address_slug = (report.address or "property").split(",")[0].replace(" ", "_")[:30]
    return GeneratedDocument(
        filename=f"LOI_{address_slug}_{now.strftime('%Y%m%d')}.txt",
        content_type="text/plain",
        data="\n".join(lines).encode("utf-8"),
    )


def generate_deal_summary(report: ZoningReport) -> GeneratedDocument:
    """Generate a Deal Summary .docx with all analysis results.

    .. deprecated::
        Use ``plotlot.clauses.engine.assemble_document()`` with
        ``DocumentType.deal_summary`` instead.
    """
    warnings.warn(
        "generate_deal_summary() is deprecated. Use plotlot.clauses.engine.assemble_document() "
        "with DocumentType.deal_summary instead.",
        DeprecationWarning,
        stacklevel=2,
    )
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return _generate_text_deal_summary(report)

    now = datetime.now(timezone.utc)
    prop = report.property_record or PropertyRecord()

    doc = Document()

    title = doc.add_heading("Deal Summary Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {now.strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    # Property Overview
    doc.add_heading("Property Overview", level=2)
    doc.add_paragraph(f"Address: {report.formatted_address or report.address}")
    doc.add_paragraph(f"Municipality: {report.municipality}")
    doc.add_paragraph(f"County: {report.county}")
    doc.add_paragraph(f"Folio: {prop.folio or 'N/A'}")
    doc.add_paragraph(f"Lot Size: {prop.lot_size_sqft:,.0f} sqft")
    doc.add_paragraph(f"Year Built: {prop.year_built or 'N/A'}")
    doc.add_paragraph("")

    # Zoning
    doc.add_heading("Zoning Analysis", level=2)
    doc.add_paragraph(f"Zoning District: {report.zoning_district}")
    doc.add_paragraph(f"Description: {report.zoning_description}")
    if report.allowed_uses:
        doc.add_paragraph(f"Allowed Uses: {', '.join(report.allowed_uses)}")
    doc.add_paragraph("")

    # Density
    if report.density_analysis:
        da = report.density_analysis
        doc.add_heading("Development Potential", level=2)
        doc.add_paragraph(f"Maximum Units: {da.max_units}")
        doc.add_paragraph(f"Governing Constraint: {da.governing_constraint}")
        doc.add_paragraph(f"Confidence: {da.confidence}")
        doc.add_paragraph("")

    # Comps
    if report.comp_analysis and report.comp_analysis.comparables:
        ca = report.comp_analysis
        doc.add_heading("Comparable Sales", level=2)
        doc.add_paragraph(f"Median Price/Acre: ${ca.median_price_per_acre:,.0f}")
        doc.add_paragraph(f"Estimated Land Value: ${ca.estimated_land_value:,.0f}")
        doc.add_paragraph(f"Confidence: {ca.confidence:.0%}")
        doc.add_paragraph(f"Comps Found: {len(ca.comparables)}")
        doc.add_paragraph("")

    # Pro Forma
    if report.pro_forma and report.pro_forma.max_land_price > 0:
        pf = report.pro_forma
        doc.add_heading("Land Pro Forma (Residual Valuation)", level=2)
        doc.add_paragraph(f"GDV: ${pf.gross_development_value:,.0f}")
        doc.add_paragraph(f"Hard Costs: ${pf.hard_costs:,.0f}")
        doc.add_paragraph(f"Soft Costs: ${pf.soft_costs:,.0f}")
        doc.add_paragraph(f"Builder Margin: ${pf.builder_margin:,.0f}")
        doc.add_paragraph(f"Max Land Price: ${pf.max_land_price:,.0f}")
        doc.add_paragraph(f"Cost per Door: ${pf.cost_per_door:,.0f}")
        doc.add_paragraph("")

    # Summary
    if report.summary:
        doc.add_heading("AI Summary", level=2)
        doc.add_paragraph(report.summary)

    # Footer
    p = doc.add_paragraph()
    p.add_run("Generated by PlotLot — ").italic = True
    p.add_run(now.strftime("%Y-%m-%d %H:%M UTC")).italic = True

    buf = io.BytesIO()
    doc.save(buf)

    address_slug = (report.address or "property").split(",")[0].replace(" ", "_")[:30]
    return GeneratedDocument(
        filename=f"DealSummary_{address_slug}_{now.strftime('%Y%m%d')}.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buf.getvalue(),
    )


def _generate_text_deal_summary(report: ZoningReport) -> GeneratedDocument:
    """Fallback plain-text deal summary."""
    now = datetime.now(timezone.utc)
    prop = report.property_record or PropertyRecord()

    lines = [
        "DEAL SUMMARY REPORT",
        f"Generated: {now.strftime('%B %d, %Y')}",
        "",
        f"Address: {report.formatted_address or report.address}",
        f"Municipality: {report.municipality}, {report.county} County",
        f"Folio: {prop.folio or 'N/A'}",
        f"Lot Size: {prop.lot_size_sqft:,.0f} sqft",
        f"Zoning: {report.zoning_district} — {report.zoning_description}",
        "",
    ]

    if report.density_analysis:
        lines.append(f"Max Units: {report.density_analysis.max_units}")
        lines.append(f"Governing Constraint: {report.density_analysis.governing_constraint}")
        lines.append("")

    if report.comp_analysis:
        ca = report.comp_analysis
        lines.append(f"Median $/Acre: ${ca.median_price_per_acre:,.0f}")
        lines.append(f"Est. Land Value: ${ca.estimated_land_value:,.0f}")
        lines.append("")

    if report.pro_forma and report.pro_forma.max_land_price > 0:
        pf = report.pro_forma
        lines.append(f"Max Land Price: ${pf.max_land_price:,.0f}")
        lines.append("")

    lines.append(f"Generated by PlotLot — {now.strftime('%Y-%m-%d %H:%M UTC')}")

    address_slug = (report.address or "property").split(",")[0].replace(" ", "_")[:30]
    return GeneratedDocument(
        filename=f"DealSummary_{address_slug}_{now.strftime('%Y%m%d')}.txt",
        content_type="text/plain",
        data="\n".join(lines).encode("utf-8"),
    )
